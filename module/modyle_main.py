import docx #docxファイルの操作
from pdf2docx.main import parse #pdf ->docxファイルに変換

import scrape   #スクレイピングをするモジュール

import time #時間
import os   #ファイル操作

with open(r"module\pdf\access_log.txt", "r")as f:
    before_access=f.readlines()[0]

before_access=[before_access[:4], before_access[4:6], before_access[6:]]
now_time=[time.strftime("%Y"), time.strftime("%m"), time.strftime("%d")]

before_access = [int(i) for i in before_access]
now_time      = [int(i) for i in now_time]

print(before_access, now_time)
if before_access[0]<now_time[0] or before_access[1]<now_time[1] or before_access[2]<now_time[2]:
    print("other day")
    #scrape.scr("https://www.ktc.ac.jp/dept/kyomu/schedule/")    #学校の行事予定のページのスクレイピング
    pdf_folder="module/pdf"
    with open(pdf_folder+"/access_log.txt", "w")as f:   #最終アクセス日の更新
        f.write(time.strftime("%Y%m%d"))

nenndo=str(int(time.strftime("%Y")) -1) if time.strftime("%m")=="01" or time.strftime("%m")=="02" or time.strftime("%m")=="03" else time.strftime("%Y")
parse("module/pdf/" + nenndo + "_yotei_zenki.pdf" , "module/pdf/" + nenndo + "_yotei_zenki.docx")   #pdf2docx
parse("module/pdf/" + nenndo + "_yotei_kouki.pdf" , "module/pdf/" + nenndo + "_yotei_kouki.docx")

with open("module/pdf/yotei.txt", "w")as f: #空白にする
    f.write("")

#要素をテキストファイルに書き出す
for i in ["zenki", "kouki"]:
    document = docx.Document("module/pdf/" + nenndo + "_yotei_"+i +".docx")
    with open("module/pdf/yotei.txt", "a")as f:
        for table1 in document.tables:
            for row1 in table1.rows:
                for cell1 in row1.cells:
                    for table2 in cell1.tables:
                        for row2 in table2.rows:
                            for cell2 in row2.cells:
                                print(cell2.text, file=f)