import docx

document = docx.Document(r"C:\Users\aki19\Downloads\output.docx")

for paragraph in document.paragraphs:
    print(paragraph.text)

with open(r"C:\Users\aki19\Downloads\output.txt", "w")as f:
    for table1 in document.tables:
        for row1 in table1.rows:
            for cell1 in row1.cells:
                for table2 in cell1.tables:
                    for row2 in table2.rows:
                        for cell2 in row2.cells:
                            print(cell2.text, file=f)

print(len(document.tables))